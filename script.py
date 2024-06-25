import os
import shutil
import datetime
import pandas as pd
import serial
import time
from openpyxl import load_workbook
from openpyxl import Workbook
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from glob import glob

# Function to parse the data blocks
def parse_data_blocks(data_lines):
    blocks = []
    current_block = []
    for line in data_lines:
        if len(current_block) == 5:
            blocks.append(current_block)
            current_block = []
        current_block.append(line.decode().strip())
    if current_block:
        while len(current_block) < 5:
            current_block.append(current_block[-1])
        blocks.append(current_block)
    return blocks

# Function to save data to the specified cells in the Excel file
def save_to_excel(blocks, excel_path):
    wb = load_workbook(excel_path)
    ws = wb.active

    cell_positions = [
        ('B3', 'D3'), ('B13', 'D13'),
        ('B23', 'D23'), ('B33', 'D33'),
        ('B43', 'D43'), ('B54', 'D54'),
        ('J54', 'L54'), ('J43', 'L43'),
        ('J33', 'L33'), ('J23', 'L23'),
        ('J13', 'L13'), ('J3', 'L3')
    ]

    for i, block in enumerate(blocks):
        if i >= len(cell_positions):
            break
        pos1, pos2 = cell_positions[i]
        start_row1 = int(pos1[1:])
        start_row2 = int(pos2[1:])
        col1 = pos1[0]
        col2 = pos2[0]
        for j, line in enumerate(block):
            if j >= 5:
                break
            data1, data2 = map(float, line.split(','))
            ws[f'{col1}{start_row1 + j}'] = data1
            ws[f'{col2}{start_row2 + j}'] = data2

        if len(block) == 4:
            data1, data2 = map(float, block[3].split(','))
            ws[f'{col1}{start_row1 + 4}'] = data1
            ws[f'{col2}{start_row2 + 4}'] = data2

    wb.save(excel_path)
    print(f"Data saved to {excel_path}")

# Function to format numbers in the dataframe
def format_number(value, decimals):
    if isinstance(value, (int, float)):
        return f"{value:.{decimals}f}"
    return value

# Function to process files
def process_files(excel_file, word_template):
    current_date = datetime.datetime.now().strftime('%Y-%m-%d')
    output_dir = os.path.join(os.getcwd(), 'output', current_date)
    os.makedirs(output_dir, exist_ok=True)

    excel_copy = os.path.join(output_dir, 'data_copy.xlsx')
    shutil.copy(excel_file, excel_copy)

    word_copy = os.path.join(output_dir, 'report_copy.docx')
    shutil.copy(word_template, word_copy)

    wb = load_workbook(excel_copy)
    sheet = wb.active
    data = []

    for row in sheet.iter_rows(min_row=77, max_row=85, min_col=5, max_col=13, values_only=True):
        data.append(row)

    df = pd.DataFrame(data)

    decimals = [2, 2, 3, 2, 2, 2, 2]
    for col, dec in zip(df.columns, decimals):
        df[col] = df[col].apply(lambda x: format_number(x, dec))

    doc = Document(word_copy)

    # Find the table with specific number of rows and columns to replace the data
    table = None
    for tbl in doc.tables:
        if len(tbl.rows) >= 9 and len(tbl.columns) >= 9:  # Assuming the table to replace has at least 9 rows and 9 columns
            table = tbl
            break

    if table is None:
        print("No suitable table found in the Word document.")
        return

    # Replace the data in the specified rows and columns
    for row_idx, row_data in enumerate(df.values):
        row_cells = table.rows[row_idx + 3].cells  # Adjust row index if necessary
        for col_idx, cell_value in enumerate(row_data):
            p = row_cells[col_idx + 4].paragraphs[0]  # Adjust column index if necessary
            p.clear()
            run = p.add_run(str(cell_value))
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.save(word_copy)
    print(f"Data saved to {excel_copy} and {word_copy}")

if __name__ == "__main__":
    current_directory = os.path.dirname(os.path.abspath(__file__))

    excel_files = glob(os.path.join(current_directory, '*.xlsx'))
    if not excel_files:
        print("No Excel file found in the directory.")
        sys.exit(1)

    excel_file = excel_files[0]
    print(f"Using Excel file: {excel_file}")

    word_files = glob(os.path.join(current_directory, '*.docx'))
    if not word_files:
        print("No Word file found in the directory.")
        sys.exit(1)

    word_template = word_files[0]
    print(f"Using Word template: {word_template}")

    # Connect to the serial port (update port name for Windows)
    ser = serial.Serial('/dev/tty.usbserial-1330', 9600, timeout=0)
    data_lines = []

    try:
        while True:
            try:
                data = ser.readline().strip()
                if data:
                    print(data)  # Debugging: print received data
                    data_lines.append(data)
                    if len(data_lines) >= 60:  # Stop after reading 60 lines of data (12 blocks of 5 lines each)
                        break
                time.sleep(1)
            except serial.SerialTimeoutException:
                print('Data could not be read')  # Handle read timeout
    except KeyboardInterrupt:
        pass

    blocks = parse_data_blocks(data_lines)

    save_to_excel(blocks, excel_file)

    process_files(excel_file, word_template)
